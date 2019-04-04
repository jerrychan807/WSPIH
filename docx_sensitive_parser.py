#!/usr/bin/env python
# -*- coding:utf-8 -*-
# author:jack
# datetime:2019-04-03 20:38
# software: PyCharm


import docx


def parser(docxfile_add, sensitive_words, result_file):
    file = docx.Document(docxfile_add)
    # file = docx.Document(r'/Users/jack/Downloads/doctest.doc') can't read doc file
    for para in file.paragraphs:
        # print(para.text）
        test_reslut = testing_senwords(para.text, sensitive_words)
        if test_reslut != None:
            with open(result_file, 'a') as f:
                f.write(test_reslut+'----'+docxfile_add+"\n")
                return

    tables = file.tables
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                test_reslut = testing_senwords(cell.text, sensitive_words)
                if test_reslut != None:
                    with open(result_file, 'a') as f:
                        f.write(test_reslut + '----' + docxfile_add+"\n")
                        return


def testing_senwords(text, sensitive_words):
    for word in sensitive_words:
        if word in text:
            return word+''


# if __name__ == "__main__":
#     parser('/Users/jack/Downloads/text.docx', ['政治面貌'], './a.txt')

